import docx
from docx import *
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_BREAK
from docx2pdf import convert
import os
from docx.shared import Inches
from docx.oxml.ns import qn, nsdecls
from docx.enum.section import WD_ORIENTATION
from docx.enum.section import WD_SECTION_START
from docx.enum.section import WD_HEADER_FOOTER_INDEX

WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"


def set_number_of_columns(section, cols):
    # column 설정해주는 함수
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(cols))


def trans():
    print("Enter the txt file name")
    txt_name = input()

    try:
        f = open(txt_name, 'rt', encoding='utf-8')
    except FileNotFoundError:
        print("Cannot find the file")
        return

    test = f.read()
    f.close()
    # 텍스트파일에 있는 내용을 받아옴
    document = Document('standard.docx')
    document._body.clear_content()
    # python-docx 사용
    # 파일 형식을 맞추기 위해 standard.docx를 불러옴
    # 원래 있던 standard파일에 내용을 없애줌

    for i in range(0, len(test) - 2):
        if test[i] == test[i + 1] == '\n':
            templist = list(test)
            templist[i + 1] = "○"
            test = "".join(templist)
            # 문단이 끝나는곳을 표시하기 위해  두번역속 개행이 나오면 "^"을 넣어서 표시했음
            # 파이썬에서 string은 수정할수 없기떄문에 list로 변환후 수정하는 과정을 거침

    test = "○" + test
    # 처음 의원이름앞에 특수기호 추가

    alist = test.split('\n')

    for i, v in enumerate(alist):
        try:
            if v[0] == "○":
                para = document.add_paragraph()
                para.paragraph_format.line_spacing = 1.2
                run = para.add_run(v)
                run.bold = True
                para.add_run("  " + alist[i + 1])
                del alist[i + 1]
                # para.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                # para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            else:
                para = document.add_paragraph()
                run = para.add_run("  " + v)
                para.paragraph_format.line_spacing = 1.2
                # para.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                # para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                # 들여쓰기 추가
        except IndexError:
            continue
    # 특수기호 ○가 문장에 앞에 있을때 글씨체 bold 로 바꿔줌

    style = document.styles['Normal']
    font = style.font
    font.name = '함초롱바탕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '함초롱바탕')
    font.size = Pt(10)
    # 글씨크기 변경 및 폰트 변경

    document.add_page_break()
    section = document.sections[0]
    set_number_of_columns(section, 2)
    # 다단 설정(layout 설정)

    filename = os.path.splitext(txt_name)[0]
    docxfilename = filename + ".docx"
    document.save(docxfilename)
    # 저장,원래 파일이름 그대로로 확장자만 바꿔서 저장

    # convert(docxfilename)
    # pdf로 변환


trans()
